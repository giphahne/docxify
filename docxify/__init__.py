import sys
import os
import argparse
from itertools import count
from itertools import repeat
from itertools import chain
from itertools import groupby

import argcomplete

from docx import Document
import base58


def punctuate(chunk):
    """
    Word-length distribution in text/prose (English?)
    """
    #return chunk + "."
    return chunk


def unpunctuate(chunk):
    return chunk.strip()


def docxify():
    """
    Entry point for application script.
    """

    description = "Docxify"
    parser = argparse.ArgumentParser(usage=None, description=description)

    parser.add_argument("-i",
                        "--input-file",
                        type=str,
                        help="File to be docxified.")

    parser.add_argument("-o", "--output-directory", type=str, help="")

    parser.add_argument("--paragraph-size",
                        type=int,
                        default=2048,
                        help=("Chunk-size of payload encoding, in bytes "
                              "(note that size of ~encoded~ paragraph will be "
                              "larger due to encoding overhead)"))

    parser.add_argument("--chapter-size",
                        type=int,
                        default=256,
                        help=("Number of ``paragraphs'' to include "
                              "in each ``chapter''"))

    parser.add_argument("-v", "--verbose", action='store_true', default=False)

    argcomplete.autocomplete(parser)
    args = parser.parse_args()

    print(f"docxifying: {args.input_file} into: {args.output_directory}",
          file=sys.stderr)

    print(f"creating directory: {args.output_directory}", file=sys.stderr)
    os.mkdir(args.output_directory)

    f = open(args.input_file, "rb")
    paragraphs = (c for c in iter(lambda: f.read(args.paragraph_size), b""))

    for chapter_number, chapter_paragraphs in groupby(zip(
            chain.from_iterable(
                (repeat(c, args.chapter_size) for c in count(1))), paragraphs),
                                                      key=lambda x: x[0]):

        document = Document()
        document.add_heading(f'Chapter {chapter_number}', 0)

        for _, paragraph in chapter_paragraphs:
            p = punctuate(base58.b58encode(paragraph).decode('utf-8'))

            if args.verbose:
                print(chapter_number, p[:30], " ... ", p[-30:])

            result = document.add_paragraph(p)

        document.save(
            os.path.join(args.output_directory,
                         f"Chapter_{chapter_number}.docx"))

    # not strictly necessary, but since file wasn't opened within a
    # control structure it makes me feel better, lol:
    f.close()


def dedocxify():
    """
    Entry point for application script.
    """

    description = "Dedocxify"
    parser = argparse.ArgumentParser(usage=None, description=description)

    parser.add_argument("-o",
                        "--output-file",
                        type=str,
                        help="File to be dedocxified.")

    parser.add_argument("-i", "--input-directory", type=str, help="")

    argcomplete.autocomplete(parser)
    args = parser.parse_args()

    print(f"dedocxifying: {args.input_directory} into: {args.output_file}",
          file=sys.stderr)

    chapters = sorted(os.listdir(args.input_directory))

    def get_docx_paragraphs(D):
        # Note that each document (aka 'chapter')
        # begins with a header paragraph:
        yield from zip(
            repeat(D),
            Document(
                docx=os.path.join(args.input_directory, D)).paragraphs[1:])

    with open(args.output_file, "wb") as f:
        for chap, p in chain.from_iterable(map(get_docx_paragraphs, chapters)):
            f.write(base58.b58decode(unpunctuate(p.text)))
